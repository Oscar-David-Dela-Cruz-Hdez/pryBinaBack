from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Manual_usuario_skill_panamericana.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr_cells[index], header, bold=True)
        set_cell_shading(hdr_cells[index], "E8EEF5")
        if widths:
            hdr_cells[index].width = Inches(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
            if widths:
                cells[index].width = Inches(widths[index])

    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Minimanual de usuario\nSkill Alexa - Distribuidora Panamericana")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Guía de uso, frases de consulta y preguntas para el cliente")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")

    document.add_paragraph()
    document.add_paragraph(
        "Este documento resume lo que puede hacer la skill, qué debe preguntar el usuario en cada parte "
        "y qué información conviene validar con el cliente antes de continuar mejorando el asistente."
    )


def build_document():
    document = Document()
    configure_document(document)
    add_cover(document)

    document.add_heading("1. Objetivo de la skill", level=1)
    document.add_paragraph(
        "La skill de Alexa funciona como un asistente administrativo para la Distribuidora Panamericana. "
        "Su propósito es permitir consultas rápidas sobre ventas, inventario y estado de pedidos, usando voz "
        "y, cuando el dispositivo lo permite, una interfaz visual APL."
    )

    document.add_heading("2. Acceso al asistente", level=1)
    document.add_paragraph(
        "Al iniciar la skill, Alexa solicita un token de administrador de cinco dígitos. Este paso evita que "
        "personas no autorizadas consulten información interna del negocio."
    )
    add_table(
        document,
        ["Forma de acceso", "Qué debe hacer el usuario"],
        [
            ["Por voz", "Decir el token de administrador. Ejemplo: mi token es uno dos tres cuatro cinco."],
            ["Por pantalla APL", "Tocar los números del teclado visual, revisar el token ingresado y tocar Ingresar."],
        ],
        widths=[1.8, 4.7],
    )

    document.add_heading("3. Menú principal", level=1)
    document.add_paragraph(
        "Después de validar el acceso, la skill muestra el menú principal. Desde ahí se puede elegir una opción "
        "por voz o por pantalla."
    )
    add_table(
        document,
        ["Opción", "Función", "Ejemplos que puede decir el usuario"],
        [
            ["Ventas", "Consulta ganancias o movimiento de mercancía.", "checa las ganancias del mes / cuánto vendimos en la semana"],
            ["Stock", "Revisa inventario general, por producto, marca o familia.", "consulta el stock general / consulta el stock de 4x4"],
            ["Pedidos", "Consulta pedidos por enviar, enviados o finalizados.", "dime los pedidos por enviar / pedidos enviados del mes"],
            ["Ayuda", "Muestra ejemplos de uso de la skill.", "ayuda / qué puedo preguntar"],
            ["Salir", "Cierra la sesión de Alexa.", "salir / detener / cancelar"],
        ],
        widths=[1.1, 2.2, 3.2],
    )

    document.add_heading("4. Consultas de ventas", level=1)
    document.add_paragraph(
        "La sección de ventas permite consultar ganancias por rango de tiempo y, en algunos casos, movimiento "
        "de mercancía por nombre de producto o familia."
    )
    add_table(
        document,
        ["Situación", "Qué debe preguntar", "Resultado esperado"],
        [
            ["Ganancias del día", "checa las ganancias del día", "Alexa responde el total de ganancias de hoy."],
            ["Ganancias de la semana", "cuánto vendimos en la semana", "Alexa responde el total de la última semana."],
            ["Ganancias del mes", "checa las ganancias del mes", "Alexa responde el total del último mes."],
            ["Rango personalizado", "de hace 15 días / dime las ganancias de hace 15 días", "Alexa calcula el total del rango indicado."],
            ["Mercancía", "checa las ventas de Barbería / ventas de 4x4 minoxidil", "Alexa busca ventas relacionadas con el producto o familia."],
        ],
        widths=[1.5, 2.6, 2.4],
    )

    document.add_heading("5. Consultas de stock", level=1)
    document.add_paragraph(
        "La sección de stock permite revisar productos con bajo inventario y buscar existencias por producto, "
        "marca o familia."
    )
    add_table(
        document,
        ["Situación", "Qué debe preguntar", "Resultado esperado"],
        [
            ["Stock general", "consulta el stock general", "Alexa indica cuántos productos tienen stock bajo."],
            ["Por producto", "consulta el stock de 4x4 minoxidil", "Alexa responde cuántas unidades quedan del producto."],
            ["Por marca", "consulta el stock de Andis / consulta el stock de 4x4", "Alexa lista productos con stock bajo de esa marca."],
            ["Por familia", "consulta el stock de Barbería", "Alexa lista productos con stock bajo de esa familia."],
        ],
        widths=[1.5, 2.8, 2.2],
    )

    document.add_heading("6. Consultas de pedidos", level=1)
    document.add_paragraph(
        "La sección de pedidos permite consultar pedidos pendientes, enviados y finalizados. Los pedidos enviados "
        "y finalizados pueden filtrarse por día, semana, mes o rango personalizado."
    )
    add_table(
        document,
        ["Situación", "Qué debe preguntar", "Resultado esperado"],
        [
            ["Por enviar", "cuántos pedidos por enviar tenemos", "Alexa cuenta pedidos pendientes o pagados."],
            ["Enviados del día", "pedidos enviados del día", "Alexa cuenta pedidos enviados hoy."],
            ["Enviados del mes", "pedidos enviados del mes", "Alexa cuenta pedidos enviados en el último mes."],
            ["Finalizados de la semana", "pedidos finalizados de la semana", "Alexa cuenta pedidos entregados en la semana."],
            ["Personalizado", "pedidos enviados de hace 15 días / pedidos finalizados de hace 15 días", "Alexa calcula el total según el rango indicado."],
        ],
        widths=[1.5, 2.8, 2.2],
    )

    document.add_heading("7. Recomendaciones de uso", level=1)
    add_bullets(
        document,
        [
            "Usar frases completas, no solo palabras sueltas, especialmente para rangos personalizados.",
            "Para rangos personalizados, decir frases como de hace 15 días o pedidos enviados de hace 15 días.",
            "Para stock, decir el nombre más claro posible del producto, marca o familia.",
            "Si Alexa no entiende, pedir ayuda o volver al menú principal.",
            "Cerrar la skill diciendo salir, detener o cancelar cuando ya no se necesite consultar información.",
        ],
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("8. Preguntas dirigidas al cliente", level=1)
    document.add_paragraph(
        "Estas preguntas sirven para conocer mejor las necesidades del cliente y validar si la skill resuelve "
        "problemas reales del negocio."
    )
    questions = [
        "¿Cuál es el principal motivo por el que usaría una skill de Alexa para consultar información administrativa?",
        "¿Qué consulta considera más importante: ventas, stock o pedidos? ¿Por qué?",
        "¿En qué momentos del día sería más útil consultar esta información por voz?",
        "¿Quiénes dentro del negocio deberían tener permiso para usar la skill?",
        "¿El token de administrador le parece suficiente como método de seguridad o preferiría otro método?",
        "¿Qué datos le gustaría ver primero al abrir la skill?",
        "¿Qué frases usaría naturalmente para preguntar por ventas, inventario o pedidos?",
        "¿Qué información de stock considera crítica: productos bajos, productos agotados, marca o familia?",
        "¿Qué tipo de reportes le gustaría agregar en una versión futura?",
        "¿Qué esperaría que haga Alexa cuando no encuentre un producto, marca, familia o pedido?",
    ]
    add_table(
        document,
        ["No.", "Pregunta para el cliente", "Respuesta / observaciones"],
        [[index + 1, question, ""] for index, question in enumerate(questions)],
        widths=[0.5, 4.2, 1.8],
    )

    document.add_heading("9. Notas para mejoras futuras", level=1)
    add_bullets(
        document,
        [
            "Agregar más frases de entrenamiento si los usuarios usan expresiones distintas a las previstas.",
            "Definir con el cliente qué datos son sensibles y quién puede consultarlos.",
            "Evaluar si el APL debe actualizarse en todas las respuestas por voz o solo en pantallas clave.",
            "Agregar reportes visuales más específicos cuando el cliente confirme qué indicadores necesita.",
        ],
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
